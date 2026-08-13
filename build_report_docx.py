# -*- coding: utf-8 -*-
"""Build the E42 drive design report (original vs optimized) as DOCX."""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\自研BMS和软件资料\E42_优化\驱动设计报告_原始vs优化.docx"

HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK = RGBColor(0x1F, 0x4D, 0x78)
INK_BLUE = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F2F4F7"
CODE_FILL = "F4F6F9"


def set_run_font(run, ascii_font="Calibri", east="宋体", size=11, bold=False,
                 color=None, italic=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east)


def set_para(par, before=0, after=6, line=1.10, align=WD_ALIGN_PARAGRAPH.LEFT):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.alignment = align


def add_body(doc, text, size=11, bold=False, color=None, after=6, line=1.10):
    p = doc.add_paragraph()
    set_para(p, after=after, line=line)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    style = doc.styles["Heading %d" % level]
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, bold=True, color=HEADING_BLUE)
        set_para(p, before=16, after=8, line=1.0)
    elif level == 2:
        set_run_font(r, size=13, bold=True, color=HEADING_BLUE)
        set_para(p, before=12, after=6, line=1.0)
    else:
        set_run_font(r, size=12, bold=True, color=HEADING_DARK)
        set_para(p, before=8, after=4, line=1.0)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, after=8, line=1.167)
        r = p.add_run(it)
        set_run_font(r, size=11)


def add_numbers(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        set_para(p, after=8, line=1.167)
        r = p.add_run(it)
        set_run_font(r, size=11)


def add_code(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        set_para(p, before=0, after=0, line=1.0)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.15)
        pf.right_indent = Inches(0.15)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_FILL)
        p._p.get_or_add_pPr().append(shd)
        r = p.add_run(ln if ln else " ")
        set_run_font(r, ascii_font="Consolas", east="宋体", size=9)
    # spacing after code block
    last = doc.paragraphs[-1]
    last.paragraph_format.space_after = Pt(6)


def set_cell_margins(table, top=80, bottom=80, start=120, end=120):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def set_table_geometry(table, widths_in, header_fill=HEADER_FILL, font_size=10):
    total_in = sum(widths_in)
    widths_dxa = [int(w * 1440) for w in widths_in]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr
    # tblW
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(total_in * 1440)))
    tblW.set(qn("w:type"), "dxa")
    # tblInd = 120
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    # borders: single grid
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tblPr.append(borders)
    # grid
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    # header row
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        if i == 0:
            trPr.append(OxmlElement("w:tblHeader"))
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[j]))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0 and header_fill:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), header_fill)
                tcPr.append(shd)
            for par in cell.paragraphs:
                set_para(par, before=0, after=0, line=1.0)
                for r in par.runs:
                    set_run_font(r, size=font_size, bold=(i == 0))


def add_table(doc, headers, rows, widths_in):
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    set_cell_margins(table)
    hdr = table.rows[0]
    for j, h in enumerate(headers):
        hdr.cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    set_table_geometry(table, widths_in)
    # spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_furniture(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    # header
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("E42 驱动系统设计报告")
    set_run_font(hr, size=8, color=MUTED)
    # footer page number
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = fp.add_run("第 ")
    set_run_font(r1, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    fp._p.append(fld)
    r2 = fp.add_run(" 页")
    set_run_font(r2, size=8, color=MUTED)


def main():
    doc = Document()
    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")
    rf.set(qn("w:eastAsia"), "宋体")
    rpr.append(rf)
    npf = normal.paragraph_format
    npf.space_after = Pt(6)
    npf.line_spacing = 1.10
    # list styles
    for style_name in ("List Bullet", "List Number"):
        st = doc.styles[style_name]
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167
    # heading styles base
    for name, sz, col in (("Heading 1", 16, HEADING_BLUE),
                          ("Heading 2", 13, HEADING_BLUE),
                          ("Heading 3", 12, HEADING_DARK)):
        st = doc.styles[name]
        st.font.size = Pt(sz)
        st.font.color.rgb = col
        st.font.bold = True

    add_page_furniture(doc)

    # Title
    p = doc.add_paragraph()
    set_para(p, before=0, after=4, line=1.0)
    r = p.add_run("E42 驱动系统设计报告")
    set_run_font(r, size=20, bold=True, color=INK_BLUE)
    p2 = doc.add_paragraph()
    set_para(p2, before=0, after=10, line=1.0)
    r2 = p2.add_run("原始模式 vs 优化模式")
    set_run_font(r2, size=13, color=MUTED)

    # Metadata
    add_body(doc, "编制日期：2026-08-12", size=9, color=MUTED, after=1, line=1.0)
    add_body(doc, "原始版：E:\\自研BMS和软件资料\\E42\\E42\\text28.slx（未改动）", size=9, color=MUTED, after=1, line=1.0)
    add_body(doc, "优化版：E:\\自研BMS和软件资料\\E42_优化\\text28.slx", size=9, color=MUTED, after=1, line=1.0)
    add_body(doc, "平台：MATLAB R2014a / Simulink，RapidECU-U2（MPC5554），驱动任务 10ms", size=9, color=MUTED, after=12, line=1.0)

    # 1
    add_heading(doc, "1  概述", 1)
    add_body(doc, "原始模型是一套基于 RapidECU 的 FSEC 赛车整车驱动控制逻辑，核心功能包括油门踏板信号处理、扭矩请求生成、制动优先、待驶（Ready-to-Drive）控制、电机 CAN 报文打包与 BMS/MCU 状态接收。原始逻辑可以正常工作，但存在三类问题：")
    add_numbers(doc, [
        "控制策略简单：扭矩请求 = 油门原始值 × 允许标志，无扰动补偿、无主动阻尼、无扭矩斜率控制；",
        "信号浪费：滑转率、车速、电机转速/扭矩等信号算出后未参与任何控制；",
        "结构与命名混乱：死代码多、块名含换行符、传递函数等命名误导。",
    ])
    add_body(doc, "优化版在不改变上电流程、任务调度、CAN 通信、RapidECU 配置等基础框架的前提下，新增 AEKF 扰动估计与主动阻尼模块，按 2026 赛事规则补充安全逻辑，并对驱动部分做了清理。")

    # 2
    add_heading(doc, "2  原始驱动模式分析", 1)
    add_heading(doc, "2.1  总体数据流", 2)
    add_code(doc, [
        "角度传感器(ch52/ch53) → 角度处理 → 油门规则 → 油门输出 → ECU → MCU_CAN发送 → 电机控制器",
        "制动与尾灯(油压) → 油门规则 刹车优先 / 制动与待驶(待驶) → 复位/使能",
        "电机测速(CAN) → 电机转速/电机扭矩（仅测量，未用于控制）",
        "前轮转速(ch6/ch7) → 滑转率计算 → 滑转率（未使用）",
    ])
    add_heading(doc, "2.2  油门踏板处理", 2)
    add_bullets(doc, [
        "双路角度传感器：角度1（ch52）、角度2（ch53），ADC 采样；",
        "分别做零位/量程标定（约 960/1800 与 970/1780）得到百分比；",
        "两路百分比差值 >10% 判不可靠，不可靠时油门数值清零；",
        "原始值经空量程/阈值等环节做范围处理。",
    ])
    add_heading(doc, "2.3  扭矩请求生成（油门规则）", 2)
    add_code(doc, [
        "油门百分比 <5%（油门死区判断） × 无刹车（刹车判断）",
        "          + 同时油门刹车冲突锁存（油门>25% 且刹车 → 0，直至松开油门恢复）",
        "          → 动力允许判断（≥1 有效）",
        "扭矩请求 = 油门原始值 × 动力允许标志",
    ])
    add_body(doc, "特点：逻辑简单可靠，但扭矩请求无斜率限制，电机阶跃响应冲击传动系；无扰动/负载补偿；无主动阻尼。")
    add_heading(doc, "2.4  待驶与安全", 2)
    add_bullets(doc, [
        "待驶激活 = 制动有效 AND 待驶按键（EV4.11.5 要求）；",
        "安全回路断开立即退出待驶（EV4.11.6）；",
        "进入待驶后鸣笛约 1.5s（EV4.12）；",
        "系统关键信号（APPS 状态）→ 安全回路输出。",
    ])
    add_heading(doc, "2.5  原始版存在的问题", 2)
    add_table(doc,
              ["问题", "说明"],
              [
                  ["滑转率未使用", "前轮转速+后轮转速算出滑转率，仅接 Terminator"],
                  ["车速/转速/扭矩未使用", "四个输出口全部接 Terminator"],
                  ["无扰动补偿", "负载突变/齿槽脉动直接反映为转速波动与负向冲击"],
                  ["无扭矩平滑", "油门阶跃直接发 CAN，传动冲击明显"],
                  ["无 CAN 失效保护", "生命计数器解码后接 Terminator，断线无兜底"],
                  ["命名混乱", "传递函数、Subsystem、youya1 等；多处块名含换行符"],
              ],
              [2.0, 4.5])

    # 3
    add_heading(doc, "3  优化后驱动模式设计", 1)
    add_heading(doc, "3.1  总体架构", 2)
    add_body(doc, "在油门规则与油门输出之间插入 Drive_AEKF 模块：")
    add_code(doc, [
        "油门规则(油门请求) ──→ Drive_AEKF ──→ 油门输出 → ECU → MCU_CAN发送",
        "                          ↑          ↑",
        "                  电机转速(CAN)  电机扭矩(CAN)",
    ])
    add_body(doc, "清理后驱动逻辑对外输出缩减为 4 个：油门输出、复位、使能、系统关键信号。")
    add_heading(doc, "3.2  AEKF 扰动估计器（降阶机械域）", 2)
    add_body(doc, "原始论文方案为 4 状态电流域 AEKF（[id,iq,ω,Td]）；按实车仅有的信号（角度传感器 + 电机转速/扭矩 CAN 反馈）降阶为 2 状态机械域模型：")
    add_code(doc, [
        "状态：x = [ω_hat; Td_hat]",
        "量测：y = ω（电机转速）",
        "输入：u = Te（电机扭矩反馈）",
        "模型：J·dω/dt = Te − Td − B·ω",
    ])
    add_body(doc, "保留论文核心机制：")
    add_bullets(doc, [
        "NIS 门控遗忘因子 λ∈[0.85, 0.98]，卡方阈值 γ=3.841（单量测自由度 1）；",
        "协方差对角元上限保护 Pcap=[100, 400]；",
        "Joseph 形式协方差更新，保证数值正定；",
        "故障安全：NIS 持续越限 20ms 前馈淡出，恢复 50ms 后投入；",
        "前馈补偿：Td_hat 经 5Hz 低通、±6 N·m 限幅后叠加；",
        "主动阻尼：指令减去 Kd·ω（Kd=0.01 N·m·s/rad）。",
    ])
    add_heading(doc, "3.3  新增安全逻辑（2026 规则）", 2)
    add_table(doc,
              ["规则", "实现"],
              [
                  ["T12.9.12 松开踏板轮上扭矩≤0", "T_base≤0 → T_opt=0，前馈不允许叠加正扭矩"],
                  ["T12.10.5 数字信号延迟≤500ms", "电机转速/扭矩连续 50 周期（500ms）不变 → 判 CAN 失效，冻结 AEKF，禁用前馈/阻尼，退化为纯油门直通"],
              ],
              [1.8, 4.7])
    add_heading(doc, "3.4  清理与命名", 2)
    add_bullets(doc, [
        "删除死代码：前轮转速、滑转率计算、车速/滑转率/电机转速/电机扭矩输出口及 Terminator、转速转轮速换算链；",
        "重命名 26 处：传递函数→油门请求、Subsystem→驱动逻辑、Subsystem2→安全回路输出、appsrule→油门规则、同时油门刹车→油门刹车冲突 等；",
        "去除所有块名中的换行符。",
    ])

    # 4
    add_heading(doc, "4  对比总结", 1)
    add_heading(doc, "4.1  控制策略对比", 2)
    add_table(doc,
              ["项目", "原始版", "优化版"],
              [
                  ["扭矩请求", "油门原始值 × 允许标志", "同上 + AEKF 前馈 − 主动阻尼 + 斜率限制"],
                  ["扰动/负载补偿", "无", "AEKF 在线估计 Td 并前馈抵消"],
                  ["主动阻尼", "无", "Kd·ω，抑制转速振荡"],
                  ["扭矩平滑", "无", "上升 2500/s、下降 5000/s（原始单位）"],
                  ["油门松开保护", "依赖原始值≈0", "显式强制 T_opt=0"],
                  ["CAN 失效保护", "无", "500ms 检测，冻结补偿"],
                  ["滑转率/轮速", "计算未用", "已删除（实车无轮速传感器）"],
                  ["待驶/鸣笛", "制动+按键、1.5s", "不变"],
                  ["制动优先", "有（软件锁存）", "不变"],
              ],
              [1.5, 2.5, 2.5])
    add_heading(doc, "4.2  接口对比", 2)
    add_table(doc,
              ["接口", "原始版", "优化版"],
              [
                  ["驱动逻辑输出", "8 个（4 个未使用）", "4 个：油门输出/复位/使能/系统关键信号"],
                  ["电机测速输出", "3 个（后轮速度未用）", "2 个：电机转速/电机扭矩"],
                  ["传感器依赖", "角度×2 + 轮速×2（实车无）", "角度×2 + 电机 CAN 反馈"],
              ],
              [1.5, 2.5, 2.5])
    add_heading(doc, "4.3  规则符合性对比", 2)
    add_table(doc,
              ["条款", "原始版", "优化版"],
              [
                  ["T12.9.5 双 APPS", "满足", "满足"],
                  ["T12.9.9 差值 10% 判不可靠", "满足", "满足"],
                  ["T12.9.12 松开踏板零扭矩", "隐式（依赖原始值）", "显式强制"],
                  ["T12.10.5 数字信号 500ms 上限", "不满足", "满足"],
                  ["EV4.11.5/4.12 待驶与鸣笛", "满足", "满足"],
              ],
              [1.8, 2.35, 2.35])
    add_heading(doc, "4.4  代码质量对比", 2)
    add_table(doc,
              ["项目", "原始版", "优化版"],
              [
                  ["死代码", "轮速/滑转率/转速/扭矩链路", "已清理"],
                  ["命名", "混乱、含换行符", "统一清晰"],
                  ["可维护性", "一般", "明显提升"],
              ],
              [1.5, 2.5, 2.5])

    # 5
    add_heading(doc, "5  验证", 1)
    add_bullets(doc, [
        "MATLAB R2014a 加载优化版：正常；",
        "Simulink update：UPDATE_OK，无结构错误；",
        "Drive_AEKF：3 输入（T_base/w_rpm/Te_raw）→ 1 输出（T_opt），接线完整；",
        "原始版 E42\\E42\\text28.slx 全程未改动（哈希校验一致）。",
    ])

    # 6
    add_heading(doc, "6  结论与后续建议", 1)
    add_body(doc, "优化版在保留原有安全逻辑与框架的前提下，解决了扭矩冲击、负载扰动、CAN 失效无保护等核心问题，并满足 2026 赛事规则对驱动控制的增量要求，同时大幅清理了死代码与命名混乱。")
    add_numbers(doc, [
        "实车标定前核对 MCU 扭矩报文缩放（当前按 0.1 N·m/LSB，RAW=10）；",
        "台架验证 AEKF 收敛与主动阻尼系数 Kd；",
        "若后续加装轮速传感器，可恢复滑转率链路并叠加牵引力控制；",
        "电机测速内 5 条未使用 CAN 解码链（主动放电、IGBT 反馈、生命计数器、电流、预充）可择机删除。",
    ])

    doc.core_properties.title = "E42 驱动系统设计报告（原始模式 vs 优化模式）"
    doc.core_properties.author = "E42"
    doc.core_properties.subject = "驱动控制设计对比报告"
    doc.save(OUT)
    print("SAVED", OUT)


if __name__ == "__main__":
    main()
