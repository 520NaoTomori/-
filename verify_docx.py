# -*- coding: utf-8 -*-
"""Structural verification of the generated DOCX."""
from docx import Document
from docx.oxml.ns import qn

path = r"E:\自研BMS和软件资料\E42_优化\驱动设计报告_原始vs优化.docx"
doc = Document(path)
print("paragraphs:", len(doc.paragraphs))
print("tables:", len(doc.tables))

heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
print("headings:", len(heads))
for h in heads:
    print("  H:", h)

ok = True
for ti, tbl in enumerate(doc.tables):
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    tblInd = tblPr.find(qn("w:tblInd"))
    grid = tbl._tbl.find(qn("w:tblGrid"))
    grid_ws = [int(g.get(qn("w:w"))) for g in grid.findall(qn("w:gridCol"))]
    row0 = tbl.rows[0]
    tc_ws = []
    for cell in row0.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        tc_ws.append(int(tcW.get(qn("w:w"))) if tcW is not None else None)
    w_total = int(tblW.get(qn("w:w"))) if tblW is not None else None
    ind = int(tblInd.get(qn("w:w"))) if tblInd is not None else None
    match = (sum(grid_ws) == w_total and tc_ws == grid_ws)
    ok = ok and match and ind == 120
    print(f"table {ti+1}: rows={len(tbl.rows)} grid={grid_ws} tcW={tc_ws} tblW={w_total} ind={ind} OK={match}")

sec = doc.sections[0]
print("margins in:", sec.top_margin.inches, sec.bottom_margin.inches,
      sec.left_margin.inches, sec.right_margin.inches)
print("ALL_TABLE_OK" if ok else "TABLE_ISSUE")
